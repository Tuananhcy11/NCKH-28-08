# -*- coding: utf-8 -*-
"""Sinh bao cao QUY TRINH XU LY DU LIEU MOI dang .docx tu ket qua thuc nghiem."""
import os, sys, datetime
# Ban Python embeddable khong tu them thu muc script vao sys.path,
# nen phai nap tuong minh: thu muc hien tai + goc src/
import os, sys
sys.path[:0] = [os.path.dirname(os.path.abspath(__file__)),
                os.path.dirname(os.path.dirname(os.path.abspath(__file__)))]
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from common import ROOT, TAB

XANH = RGBColor(0x1F, 0x3A, 0x5F)
XAM = RGBColor(0x55, 0x55, 0x55)
DO = RGBColor(0xA3, 0x1D, 0x1D)
LUC = RGBColor(0x1B, 0x6B, 0x3A)


# ----------------------------------------------------------------- tien ich
def dat_font_chuan(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15


def vien_duoi(p, mau="1F3A5F", day=8):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(day))
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), mau)
    b.append(bot)
    pPr.append(b)


def to_mau(o, mau):
    tcPr = o._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), mau)
    tcPr.append(sh)


def bang(doc, tieu_de, hang, rong=None, canh_phai=None, cao_sang=None):
    """Tao bang co dinh dang thong nhat. cao_sang: tap chi so hang can to dam."""
    t = doc.add_table(rows=1, cols=len(tieu_de))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, ten in enumerate(tieu_de):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(str(ten))
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        to_mau(hdr[i], "1F3A5F")
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for k, h in enumerate(hang):
        o = t.add_row().cells
        for i, v in enumerate(h):
            o[i].text = ""
            p = o[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(9.5)
            if cao_sang and k in cao_sang:
                r.bold = True
            if canh_phai and i in canh_phai:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if k % 2 == 1:
            for c in o:
                to_mau(c, "F2F5F9")
    if rong:
        assert sum(rong) <= 16.65, "Tong do rong cot %.2f cm vuot be rong trang" % sum(rong)
        t.autofit = False
        for row in t.rows:
            for i, w in enumerate(rong):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def h(doc, muc, chu):
    p = doc.add_heading(chu, level=muc)
    for r in p.runs:
        r.font.color.rgb = XANH
        r.font.name = "Calibri"
    return p


def doan(doc, chu, dam=False, mau=None, co=11, sau=6):
    p = doc.add_paragraph()
    r = p.add_run(chu)
    r.bold = dam
    r.font.size = Pt(co)
    if mau:
        r.font.color.rgb = mau
    p.paragraph_format.space_after = Pt(sau)
    return p


def gach_dau(doc, chu, muc=0):
    p = doc.add_paragraph(chu, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + 0.6 * muc)
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def khung_canh_bao(doc, nhan, chu, mau_nen="FDECEA", mau_chu=DO):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    o = t.rows[0].cells[0]
    to_mau(o, mau_nen)
    o.text = ""
    p = o.paragraphs[0]
    r = p.add_run(nhan + "  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = mau_chu
    r2 = p.add_run(chu)
    r2.font.size = Pt(10)
    doc.add_paragraph()
    return t


# ------------------------------------------------------------------ noi dung
def tao(duong_dan):
    doc = Document()
    dat_font_chuan(doc)
    s = doc.sections[0]
    s.page_width = Cm(21.0)          # A4 - mac dinh cua python-docx la Letter
    s.page_height = Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = s.bottom_margin = Cm(2.0)
    # be rong dung duoc = 21.0 - 2*2.2 = 16.6 cm; tong cot moi bang phai <= 16.6

    # ---------------------------------------------------------- trang bia
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = XANH
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KHOA CÔNG NGHỆ TÀI CHÍNH"); r.font.size = Pt(12); r.font.color.rgb = XAM
    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QUY TRÌNH XỬ LÝ DỮ LIỆU MỚI")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = XANH
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Word khong hieu ky tu xuong dong trong mot run - phai dung add_break()
    r = p.add_run("Tái tạo phiên cuối tuần của XAU/USD từ PAXG")
    r.font.size = Pt(14); r.font.color.rgb = XAM
    r.add_break()
    r2 = p.add_run("và khung kiểm chứng bắt buộc")
    r2.font.size = Pt(14); r2.font.color.rgb = XAM
    vien_duoi(p)
    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bản hiệu chỉnh sau kiểm định thực nghiệm trên 279–319 kỳ nghỉ cuối tuần")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = XAM
    r.add_break()
    r2 = p.add_run("giai đoạn 2020-08 → 2025-12")
    r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = XAM

    for _ in range(6):
        doc.add_paragraph()
    bang(doc,
         ["Hạng mục", "Nội dung"],
         [["Phạm vi", "Quy trình xử lý dữ liệu, không bao gồm phần mô hình hóa giao dịch"],
          ["Thay thế cho", "Bước 4 – Bước 6 của quy trình mười bước hiện hành"],
          ["Căn cứ hiệu chỉnh", "Ba thí nghiệm kiểm chứng độc lập, mã nguồn kèm theo"],
          ["Dữ liệu kiểm định", "279–319 kỳ nghỉ cuối tuần thật, ba khung thời gian"],
          ["Ngày ban hành", datetime.date.today().strftime("%d/%m/%Y")]],
         rong=[4.5, 12.0])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------ tom tat dieu hanh
    h(doc, 1, "TÓM TẮT ĐIỀU HÀNH")
    doan(doc, "Quy trình cũ (Bước 4–6) đã được kiểm định trên dữ liệu có đáp án và cho kết quả "
              "không dùng được. Tài liệu này trình bày quy trình thay thế, cùng toàn bộ bằng chứng "
              "định lượng dẫn tới từng thay đổi.")

    khung_canh_bao(doc, "PHÁT HIỆN CHÍNH",
                   "Quy trình cũ tạo ra sai số 32,54 USD khi ước lượng giá mở cửa phiên đầu tuần, "
                   "trong khi giả định đơn giản “giá không đổi” chỉ sai 2,84 USD. "
                   "Nghĩa là toàn bộ bộ máy tái tạo làm kết quả xấu đi gấp 11,5 lần so với việc không làm gì.")

    doan(doc, "Ba nguyên nhân gốc", dam=True, co=12)
    gach_dau(doc, "Hệ số beta lấy từ trong phiên (0,86–0,93) trong khi beta thật của cuối tuần "
                  "chỉ là 0,15–0,23 — sai lệch 4 đến 6 lần.")
    gach_dau(doc, "Số hạng nhiễu ε được cộng vào mỗi nến, tích lũy theo bước ngẫu nhiên qua 50–190 "
                  "nến cuối tuần, tạo độ trôi khoảng 53 USD.")
    gach_dau(doc, "Hàm mục tiêu hiệu chỉnh ở Bước 5 nhắm vào tương quan, một đại lượng trực giao "
                  "với độ chính xác mức giá.")

    doan(doc, "Bốn thay đổi bắt buộc", dam=True, co=12)
    bang(doc,
         ["#", "Thay đổi", "Lý do định lượng"],
         [["1", "Neo theo lợi suất, không neo theo mức giá PAXG",
           "Tránh thừa hưởng sai lệch tỷ lệ delta ~38 bp mỗi kỳ nghỉ"],
          ["2", "Dùng beta cuối tuần ước lượng cuộn chiếu (0,15 / 0,087)",
           "Beta trong phiên làm sai số tăng 136–172 %"],
          ["3", "Bỏ hoàn toàn số hạng nhiễu ε khi dựng đặc trưng",
           "ε tích lũy tạo sai số 30–32 USD trên mức giá"],
          ["4", "Đổi hàm mục tiêu hiệu chỉnh sang sai số dự báo",
           "Khớp tương quan không bảo đảm mức giá đúng"]],
         rong=[1.0, 7.0, 8.5], cao_sang={1})

    doan(doc, "Kỳ vọng thực tế sau hiệu chỉnh: sai số khoảng 2,7 USD so với 2,84 USD của mốc "
              "“không làm gì” — cải thiện 3–4 %, với trần lý thuyết là 12,8 %. Giá trị khai thác "
              "được từ PAXG cho phiên cuối tuần là nhỏ và có giới hạn cứng, bất kể phương pháp nào.",
         mau=XAM)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------- phan I: bang chung
    h(doc, 1, "PHẦN I — BẰNG CHỨNG DẪN TỚI HIỆU CHỈNH")

    h(doc, 2, "1.1. Thí nghiệm 1: Tương quan qua kỳ nghỉ so với trong phiên")
    doan(doc, "Đo trực tiếp trên 279–319 kỳ nghỉ thật, so sánh với hệ số nền của Bước 3.")
    bang(doc,
         ["Đại lượng", "D1", "H1", "M15"],
         [["Biến động cuối tuần của XAU (|r| trung bình)", "12,9 bp", "12,7 bp", "11,8 bp"],
          ["Biến động cuối tuần của PAXG (|r| trung bình)", "35,1 bp", "37,3 bp", "34,4 bp"],
          ["Tỷ lệ PAXG / XAU", "2,7×", "2,9×", "2,9×"],
          ["Tương quan QUA KỲ NGHỈ", "0,489", "0,339", "0,336"],
          ["Tương quan TRONG PHIÊN (Bước 3)", "0,905", "0,663", "0,463"],
          ["Beta QUA KỲ NGHỈ", "0,231", "0,147", "0,145"],
          ["Beta TRONG PHIÊN (Bước 3)", "0,897", "0,579", "0,350"]],
         rong=[8.0, 2.8, 2.8, 2.8], canh_phai=[1, 2, 3], cao_sang={3, 5})
    doan(doc, "PAXG dao động gấp gần ba lần vàng trong kỳ nghỉ, và phần dao động thêm là nhiễu riêng "
              "của thị trường tiền mã hóa cuối tuần. Quan hệ chênh lệch giá cần thị trường giao ngay "
              "mở cửa mới hoạt động; cuối tuần nó bị ngắt.")

    h(doc, 2, "1.2. Thí nghiệm 2: Suy giảm độ chính xác khi thị trường đóng cửa")
    doan(doc, "Đóng băng nhóm đặc trưng “trí nhớ Delta” đúng như khi không có nến XAU nào, "
              "chạy trên 400 khối × 50 nến khung H1. MAE thành phần close, đơn vị USD.")
    bang(doc,
         ["Chân trời", "Một bước (trần trên)", "Mô hình đóng băng", "Giữ delta cố định"],
         [["1 nến", "2,33", "2,57", "2,50"],
          ["2–4 nến", "2,44", "3,42", "3,52"],
          ["5–12 nến", "2,40", "4,12", "4,51"],
          ["13–24 nến", "2,49", "4,78", "5,22"],
          ["trên 24 nến", "2,54", "5,72", "6,37"]],
         rong=[4.0, 4.5, 4.0, 4.0], canh_phai=[1, 2, 3], cao_sang={4})
    doan(doc, "Sai số tăng 2,2 lần khi kéo dài chân trời, trong khi trần trên đứng yên ở 2,5 USD. "
              "Toàn bộ phần tăng thêm đến từ việc mất trí nhớ, không phải từ độ khó của chuỗi. "
              "R² tương ứng tụt từ 0,925 xuống 0,675.")

    h(doc, 2, "1.3. Thí nghiệm 3: So sánh sáu cách xử lý trên cùng bộ kỳ nghỉ")
    doan(doc, "MAE ước lượng giá mở cửa phiên đầu tuần, đơn vị USD.")
    bang(doc,
         ["Xếp hạng", "Cách xử lý", "D1", "H1", "M15", "So với “không đổi”"],
         [["1", "Giả định giá không đổi (không dùng PAXG)", "2,81", "2,84", "2,61", "mốc"],
          ["2", "Neo lợi suất với beta = 0,15", "2,72", "2,93", "2,70", "+3,5 %"],
          ["3", "Pipeline delta OHLC (hồi quy cuộn chiếu)", "8,45", "3,91", "3,22", "+38 %"],
          ["4", "Neo lợi suất với beta = 0,90", "6,71", "7,80", "7,15", "+175 %"],
          ["5", "Giữ nguyên delta (beta = 1)", "7,42", "8,56", "7,85", "+202 %"],
          ["6", "QUY TRÌNH CŨ đúng như đã triển khai", "30,99", "32,54", "30,62", "+1 048 %"]],
         rong=[1.6, 6.0, 2.0, 2.0, 2.0, 3.0], canh_phai=[2, 3, 4, 5], cao_sang={0, 5})

    khung_canh_bao(doc, "CƠ CHẾ HỎNG",
                   "σ_e ≈ 22 bp mỗi nến × κ = 1,5437 → 34 bp mỗi nến × √50 nến cuối tuần "
                   "→ 240 bp ≈ 2,4 % ≈ 53 USD. Khớp với sai số đo được 32–43 USD. "
                   "Đây là lỗi thiết kế hàm mục tiêu, không phải lỗi lập trình.",
                   mau_nen="FFF6E5", mau_chu=RGBColor(0x8A, 0x5A, 0x00))

    h(doc, 2, "1.4. Thí nghiệm 4: Chất lượng đặc trưng biến động")
    doan(doc, "Quy trình cũ không nhằm tạo mức giá mà tạo ba đặc trưng. Đo ở đúng mục đích đó: "
              "tương quan giữa biến động cuối tuần ước lượng và biến động thật của vàng.")
    bang(doc,
         ["Đặc trưng", "D1", "H1", "M15"],
         [["Chuỗi 24/7 của quy trình cũ (có beta, tin tức, ε)", "0,241", "0,092", "0,148"],
          ["|r_PAXG| thô — không beta, không ε, không tin tức", "0,346", "0,169", "0,183"]],
         rong=[9.0, 2.5, 2.5, 2.5], canh_phai=[1, 2, 3], cao_sang={1})
    doan(doc, "Toàn bộ bộ máy phân tầng beta, điều biến tin tức và vòng lặp hiệu chỉnh κ làm đặc trưng "
              "biến động XẤU ĐI so với việc chỉ lấy trị tuyệt đối lợi suất PAXG. Ở khung H1, tương quan "
              "giảm gần một nửa. Đây là lời giải thích cho kết quả âm tính ở Bước 10: ba đặc trưng cuối "
              "tuần xếp hạng 23–37 trên 40 theo SHAP vì chúng được dựng từ một chuỗi đã bị làm nhiễu "
              "có chủ đích.", mau=XAM)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --------------------------------------------- phan II: quy trinh moi
    h(doc, 1, "PHẦN II — QUY TRÌNH XỬ LÝ MỚI")

    h(doc, 2, "2.1. Sơ đồ tổng thể")
    bang(doc,
         ["Giai đoạn", "Nội dung", "Đầu ra"],
         [["A. Thu thập", "Mỗi chuỗi MỘT nguồn duy nhất, không ghép sàn",
           "xau_{tf}, paxg_{tf}, usdtusd_{tf}, gdelt"],
          ["B. Làm sạch", "Khử de-peg USDT, dựng lại nến PAXG trên lưới XAU",
           "df_aligned"],
          ["C. Kiểm chứng nguồn", "Tương quan chéo có độ trễ để bắt lệch múi giờ",
           "báo cáo lag = 0"],
          ["D. Ước lượng beta", "Beta CUỐI TUẦN cuộn chiếu, không dùng beta trong phiên",
           "beta_cuoi_tuan_{tf}"],
          ["E. Tái tạo", "Neo lợi suất từ giá đóng thứ Sáu, KHÔNG cộng nhiễu",
           "chuoi_247_{tf}"],
          ["F. Trích đặc trưng", "Ba đặc trưng gắn vào nến mở phiên đầu tuần",
           "dac_trung_cuoi_tuan_{tf}"],
          ["G. Kiểm định", "Đối chiếu mốc “không đổi”, ADF/KPSS, lookahead",
           "bảng nghiệm thu"]],
         rong=[3.2, 8.0, 5.0])

    h(doc, 2, "2.2. Giai đoạn A — Thu thập, nguyên tắc một nguồn")
    doan(doc, "Ghép hai sàn vào cùng một chuỗi giá tạo gãy cấu trúc giả tại điểm nối: chênh lệch "
              "mua-bán, giờ chốt phiên và quy ước làm tròn khác nhau khiến log return tại mốc nối "
              "là giả tạo. Bằng chứng đo được: khi bỏ ghép nguồn cho chuỗi USDT/USD, lệch neo cực "
              "đại giảm từ 280 bp xuống 39–58 bp.")
    bang(doc,
         ["Chuỗi", "Nguồn duy nhất", "Phạm vi thực tế"],
         [["XAU/USD M15 / H1 / D1", "Dukascopy qua data.forexsb.com",
           "H1, D1: 2015-01 → 2025-12; M15: 2018-03 →"],
          ["PAXG/USDT", "Binance REST klines", "2020-08-28 →"],
          ["USDT/USD", "Bitfinex tUSTUSD", "2018-11-27 →"],
          ["XAUT/USD (kiểm chứng chéo)", "Bitfinex tXAUT:USD", "2020-01-24 →"],
          ["Tin tức", "GDELT DOC 2.0, sáu chuỗi", "2017-01 →"]],
         rong=[4.5, 6.0, 6.0])
    khung_canh_bao(doc, "BẮT BUỘC",
                   "Nến D1 phải gộp theo phiên ngoại hối 22:00 → 22:00 UTC. Nếu gộp theo mốc 00:00 UTC, "
                   "phiên mở cửa tối Chủ nhật trở thành một nến D1 giả chỉ dài 2 giờ — lỗi này từng "
                   "sinh ra 565 nến giả, đẩy số nến từ 258 lên 311 mỗi năm.",
                   mau_nen="FFF6E5", mau_chu=RGBColor(0x8A, 0x5A, 0x00))

    h(doc, 2, "2.3. Giai đoạn B — Làm sạch và đồng bộ")
    doan(doc, "Khử de-peg USDT:", dam=True)
    doan(doc, "        P_PAXG_USD,col(t) = P_PAXG_USDT,col(t) × E_USDT/USD(t)")
    doan(doc, "        peg_dev(t) = | E_USDT/USD(t) − 1,0 |")
    doan(doc, "Tỷ giá neo lấy từ sàn fiat độc lập với sàn báo giá PAXG, ghép LÙI để chỉ dùng thông tin "
              "đã công bố tại hoặc trước mốc mở nến. Độ phủ oracle đạt 100 %.")
    doan(doc, "Đồng bộ mốc thời gian:", dam=True)
    doan(doc, "Nến D1 của XAU dán nhãn tại 22:00 UTC còn nến D1 của Binance tại 00:00 UTC — ghép trực "
              "tiếp cho 0 mốc trùng. Khi tỷ lệ trùng lưới dưới 50 %, phải dựng lại nến PAXG trên đúng "
              "lưới nến XAU: với mỗi nến XAU bắt đầu tại t, gom nến PAXG trong [t, t + độ_dài) rồi tính "
              "lại OHLC. Giới hạn cứng độ_dài là bắt buộc, nếu không dữ liệu cuối tuần sẽ bị hút vào "
              "nến đầu tuần.")

    h(doc, 2, "2.4. Giai đoạn C — Kiểm chứng nguồn (bước mới, bắt buộc)")
    doan(doc, "Mỗi khi đổi nguồn dữ liệu, phải chạy tương quan chéo có độ trễ giữa log return XAU và "
              "PAXG. Đỉnh phải nằm đúng tại độ trễ 0. Một sai lệch 1–2 giờ sẽ phá hủy toàn bộ kết quả "
              "về sau mà không để lại dấu hiệu rõ ràng.")
    bang(doc,
         ["Độ trễ", "H1", "M15", "Diễn giải"],
         [["−2", "0,018", "0,033", ""],
          ["−1", "0,074", "0,104", ""],
          ["0", "0,658", "0,467", "đỉnh đúng vị trí — dấu thời gian chuẩn UTC"],
          ["+1", "−0,003", "−0,003", ""],
          ["+2", "−0,004", "0,006", ""]],
         rong=[2.2, 2.5, 2.5, 9.0], canh_phai=[1, 2], cao_sang={2})

    h(doc, 2, "2.5. Giai đoạn D — Ước lượng beta cuối tuần (thay đổi cốt lõi)")
    khung_canh_bao(doc, "THAY ĐỔI SO VỚI QUY TRÌNH CŨ",
                   "Không dùng beta của Bước 3. Beta trong phiên đo quan hệ khi cả hai thị trường "
                   "cùng mở; beta cuối tuần là một tham số KHÁC và phải ước lượng riêng.")
    doan(doc, "Quy tắc ước lượng:", dam=True)
    gach_dau(doc, "Chỉ dùng các kỳ nghỉ ĐÃ QUA (ước lượng cuộn chiếu), không nhìn trước.")
    gach_dau(doc, "Cần tối thiểu 30 kỳ nghỉ trước khi bắt đầu áp dụng.")
    gach_dau(doc, "Hồi quy: r_XAU_kỳ_nghỉ = β · r_PAXG_kỳ_nghỉ + sai số.")
    gach_dau(doc, "Giá trị hội tụ: β ≈ 0,176 ở D1 và ≈ 0,087 ở H1 / M15.")
    doan(doc, "Kiểm chứng đối sánh trực tiếp (MAE giá mở cửa đầu tuần, USD):")
    bang(doc,
         ["Cách chọn beta", "D1", "H1", "M15", "So với “không đổi”"],
         [["Giả định giá không đổi (β = 0)", "2,875", "2,899", "2,639", "mốc"],
          ["Beta trong phiên = 0,90 (quy trình cũ)", "6,791", "7,790", "7,166", "+136 % đến +172 %"],
          ["Beta cuối tuần cuộn chiếu (quy trình mới)", "2,754", "2,793", "2,544", "−3,6 % đến −4,2 %"]],
         rong=[6.0, 2.2, 2.2, 2.2, 4.0], canh_phai=[1, 2, 3], cao_sang={2})

    h(doc, 2, "2.6. Giai đoạn E — Tái tạo chuỗi cuối tuần")
    doan(doc, "Công thức mới:", dam=True)
    doan(doc, "        r̂_t = β_cuối_tuần × r_PAXG,t")
    doan(doc, "        XAU_ước_lượng(T) = XAU_đóng_thứ_Sáu × exp( Σ r̂_t )")
    doan(doc, "So với công thức cũ:", dam=True)
    doan(doc, "        r̂_t = β_trong_phiên × s × r_PAXG,t × (1 + λ·I_t) + ε_t        ← ĐÃ LOẠI BỎ")
    bang(doc,
         ["Thành phần", "Quy trình cũ", "Quy trình mới", "Lý do"],
         [["Hệ số beta", "0,86–0,93 (trong phiên)", "0,087–0,176 (cuối tuần)",
           "Beta trong phiên sai 4–6 lần"],
          ["Điều biến tin tức (1 + λI)", "Có, λ = 0,30", "Bỏ",
           "Làm đặc trưng biến động xấu đi"],
          ["Nhiễu đặc thù ε", "Có, κ = 1,5437", "Bỏ",
           "Tích lũy tạo sai số 30–32 USD"],
          ["Neo vào", "Giá đóng thứ Sáu", "Giữ nguyên",
           "Đúng, tránh sai lệch mức giá PAXG"],
          ["Vòng lặp hiệu chỉnh", "Khớp tương quan ±0,05", "Cực tiểu sai số dự báo",
           "Mục tiêu cũ trực giao với độ chính xác"]],
         rong=[3.8, 4.0, 4.0, 4.7], cao_sang={0, 2})

    khung_canh_bao(doc, "GIỮ ε KHI NÀO",
                   "Chỉ giữ số hạng ε khi cần mô phỏng Monte Carlo phân phối kết quả. Khi đó phải "
                   "sinh nhiều đường mẫu và dùng phân vị, tuyệt đối không lấy một đường mẫu duy nhất "
                   "làm dữ liệu đầu vào cho mô hình học máy.",
                   mau_nen="EAF4EC", mau_chu=LUC)

    h(doc, 2, "2.7. Giai đoạn F — Trích đặc trưng")
    doan(doc, "Nguyên tắc bất di bất dịch giữ nguyên từ quy trình cũ: chuỗi 24/7 KHÔNG dùng để sinh "
              "lệnh. Mọi giao dịch và chỉ tiêu hiệu suất tính trên chuỗi 24/5 gốc.")
    bang(doc,
         ["#", "Đặc trưng", "Công thức mới", "Thay đổi"],
         [["1", "Biến động cuối tuần", "|r_PAXG| thô qua kỳ nghỉ",
           "Thay cho √(Σ r̂²) — tương quan cao hơn 1,4–1,8 lần"],
          ["2", "Độ lệch tích lũy PAXG–XAU", "Σ r_PAXG − Σ r̂ (với β mới)",
           "Giữ cấu trúc, đổi β"],
          ["3", "Điểm tâm lý cuối tuần", "Trung bình tone GDELT trong kỳ nghỉ",
           "Giữ nguyên"]],
         rong=[1.0, 4.5, 5.5, 5.5], cao_sang={0})
    doan(doc, "Mỗi kỳ nghỉ gán vào cây nến 24/5 đầu tiên sau đó. Các nến khác nhận giá trị 0 kèm cờ "
              "nhị phân báo “nến này có/không có thông tin cuối tuần”, tránh để mô hình hiểu nhầm "
              "0 là một quan sát thật.")

    h(doc, 2, "2.8. Giai đoạn G — Khung nghiệm thu bắt buộc")
    doan(doc, "Không chấp nhận bộ dữ liệu nào chưa qua đủ năm phép kiểm dưới đây.")
    bang(doc,
         ["#", "Phép kiểm", "Tiêu chí đạt"],
         [["1", "Đối chiếu mốc “giá không đổi”",
           "MAE của phương pháp phải THẤP HƠN mốc này, nếu không thì loại bỏ"],
          ["2", "Đối chiếu mô hình ngây thơ",
           "R² phải vượt R² của dự báo Δ̂ₜ = Δₜ₋₁, và phải báo cáo cả hai"],
          ["3", "Kiểm tra lookahead tự động",
           "Không đặc trưng nào đạt đỉnh tương quan ở độ trễ dương"],
          ["4", "ADF và KPSS trên phần dư",
           "Hai kiểm định phải nhất trí kết luận phần dư dừng"],
          ["5", "Ràng buộc hình học OHLC",
           "Báo cáo tỷ lệ vi phạm trước khi kẹp; sau khi kẹp phải bằng 0"]],
         rong=[1.0, 5.5, 10.0])

    khung_canh_bao(doc, "QUY TẮC BÁO CÁO",
                   "Khi trình bày chuỗi 24/7, bắt buộc khai báo sai số ±2,8 USD. Không được trích các "
                   "con số R² 0,94–0,97 của dự báo một bước như thể chúng mô tả chất lượng dữ liệu "
                   "cuối tuần — hai đại lượng đo hai điều kiện hoàn toàn khác nhau.")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # -------------------------------------------------- phan III: trien khai
    h(doc, 1, "PHẦN III — TRIỂN KHAI VÀ GIỚI HẠN")

    h(doc, 2, "3.1. Bản đồ mã nguồn")
    bang(doc,
         ["Tệp", "Giai đoạn", "Chức năng"],
         [["step01e_forexsb.py", "A", "XAU/USD ba khung từ Dukascopy"],
          ["step01_collect.py", "A", "PAXG/USDT; GLD kiểm chứng chéo"],
          ["step01c_bitfinex.py", "A", "USDT/USD và XAUT/USD"],
          ["step01b_gdelt.py", "A", "Sáu chuỗi tin GDELT"],
          ["delta_ohlc/preprocessing.py", "B", "Khử de-peg, dựng lại nến trên lưới XAU"],
          ["delta_ohlc/features.py", "C, F", "Ma trận đặc trưng, dò lookahead"],
          ["delta_ohlc/model.py", "D", "Hồi quy trượt đa mục tiêu"],
          ["delta_ohlc/evaluate.py", "G", "MAE/RMSE/R², ADF, KPSS"],
          ["run_delta_gap_test.py", "G", "Kiểm định trên kỳ nghỉ thật"],
          ["run_so_sanh_hai_phuong_phap.py", "G", "Đối chiếu sáu cách xử lý"]],
         rong=[6.5, 2.5, 7.5])

    h(doc, 2, "3.2. Ba giới hạn phải khai báo")
    bang(doc,
         ["#", "Giới hạn", "Ảnh hưởng", "Khắc phục"],
         [["1", "PAXG không tồn tại trước 09/2019; Binance niêm yết 08/2020",
           "Beta cuối tuần chỉ ước lượng được từ 2020-08",
           "Không khắc phục được — giới hạn của tài sản"],
          ["2", "XAU M15 chỉ từ 2018-03 do giới hạn 200 000 nến mỗi tệp",
           "Khung M15 dựa trên 7,8 năm thay vì 11 năm",
           "API key Twelve Data / Polygon, hoặc xuất M1 từ MetaTrader 5"],
          ["3", "FinBERT chưa dùng đúng nghĩa, thay bằng tone GDELT tổng hợp",
           "Đặc trưng tâm lý yếu hơn thiết kế",
           "Tải artlist rồi chấm điểm từng tiêu đề"]],
         rong=[1.0, 5.0, 5.0, 5.5])

    h(doc, 2, "3.3. Kết luận")
    doan(doc, "Quy trình mới sửa được ba lỗi thiết kế nghiêm trọng của quy trình cũ và đưa sai số từ "
              "32,54 USD xuống khoảng 2,7 USD. Tuy nhiên phải trung thực rằng mức 2,7 USD này chỉ tốt "
              "hơn 3–4 % so với việc giả định giá không đổi, và trần lý thuyết đo được từ R² qua kỳ "
              "nghỉ chỉ là 12,8 %.")
    doan(doc, "Khuyến nghị sử dụng:", dam=True)
    gach_dau(doc, "Dùng chuỗi tái tạo cho BIẾN ĐỘNG, không dùng cho HƯỚNG. Tương quan 0,34–0,49 quá "
                  "thấp để suy ra chiều đi của giá.")
    gach_dau(doc, "Trình bày kết quả âm tính này như một đóng góp học thuật độc lập: nó trả lời dứt "
                  "khoát câu hỏi PAXG có thay thế được vàng giao ngay ngoài giờ hay không.")
    gach_dau(doc, "Giữ toàn bộ khung nghiệm thu năm phép kiểm cho mọi nghiên cứu về sau.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    vien_duoi(p, "CCCCCC", 4)
    doan(doc, "Toàn bộ số liệu trong tài liệu này sinh từ mã nguồn kèm theo, hạt giống ngẫu nhiên cố "
              "định SEED = 42, có thể tái lập.", mau=XAM, co=9)

    doc.save(duong_dan)
    return duong_dan


if __name__ == "__main__":
    out = os.path.join(ROOT, "Quy_Trinh_Xu_Ly_Du_Lieu_Moi.docx")
    tao(out)
    print("Da tao:", out, os.path.getsize(out), "byte")
